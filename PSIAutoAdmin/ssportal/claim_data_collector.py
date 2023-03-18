from playwright.sync_api import sync_playwright
from PSIAutoAdmin.util import csv_processor,browser_initiator
from datetime import datetime, time
from pathlib import Path
from docx import Document
from docx.shared import Cm
import re
import json
import logging

MODULE_LOGGER = logging.getLogger()

def append_data_document(data_csv,document):
    """
    Method to compile all claimable date and corresponding image to docx.
    Needed to be uploaded in bipo claim process.
    """
    MODULE_LOGGER.debug("Docx Append method called")
    for individual_data in data_csv:
        MODULE_LOGGER.info(f"Screenshot for date {individual_data['docx_date']} will appended to docx")
        document.add_paragraph(individual_data["docx_date"])
        document.add_picture(individual_data["docx_image_path"],width=Cm(22.86))


def claimable_data_collector(page,screenshot_folder):
    """
    The main actual data processing method after the site is navigated.
    Will return & generate data for both CSV & DOCX file needed by bipo claim process automation.
    """
    MODULE_LOGGER.debug("Claimable data collector method is called")
    # Wait for table to be shown
    MODULE_LOGGER.debug("Waiting for table")
    page.wait_for_load_state(state="networkidle")
    page.locator("table#custList").wait_for()

    # Select option to list 100 row
    MODULE_LOGGER.debug("List latest 100 row")
    page.locator("tbody >> tr >> td >> select").select_option("100")
    page.locator("tr >> internal:has=\"td:text('100')\"").wait_for()

    # Collect approved row only
    row = page.locator("tr >> internal:has=\"td:text('Approved')\"")

    # Initiate list for final data
    csv_data_list=[]    
    document_data_list=[]

    # Last claim date to filter and skip previous claimed date
    try:
        with open('config.json', 'r') as stream:
            data_json = json.load(stream)
            config_data = data_json["ssportal_claim"]
        tanggal_claim_terakhir=datetime.strptime(config_data['last_claimed'],'%d %b %Y')
        MODULE_LOGGER.info(f"Will skip data prior to {config_data['last_claimed']}")
    except:
        tanggal_claim_terakhir=datetime(datetime.date.today().year,1,1)
        MODULE_LOGGER.warning(f"json not found, will try to skip data prior 1st day of this year")

    for index in range(row.count()):
        # Collect data from ssportal table row
        customer = row.nth(index).locator("role=gridcell").nth(7).text_content()
        project = row.nth(index).locator("role=gridcell").nth(9).text_content()
        so = row.nth(index).locator("role=gridcell").nth(10).text_content()
        activity_description = row.nth(index).locator("role=gridcell").nth(4).text_content()
        tanggal_str = row.nth(index).locator("role=gridcell").nth(1).text_content()
        start_hour_str = row.nth(index).locator("role=gridcell").nth(2).text_content()
        end_hour_str = row.nth(index).locator("role=gridcell").nth(3).text_content()

        # Convert date and time string to date/time object for comparison purpose
        tanggal = datetime.strptime(tanggal_str,'%d %b %Y')
        start_hour = datetime.strptime(start_hour_str, "%H:%M").time()
        end_hour = datetime.strptime(end_hour_str, "%H:%M").time()

        screenshot_naming = f"ssportal_claim_data - {tanggal.strftime('%d %B %Y')}.png"
        
        if tanggal_claim_terakhir < tanggal:
            MODULE_LOGGER.info(f'Processing data number {index+1}')
            print(f'Processing data number {index+1}')
            # Compare data to get nominal for eligible claim
            if tanggal.date().weekday() < 5:
                if start_hour >= time(20) or end_hour >= time(20):
                    nominal = 50000
                elif start_hour >= time(0) and start_hour < time(4) or end_hour >= time(0) and end_hour < time(8):
                    nominal = 100000
                else:
                    nominal = 0
            elif tanggal.date().weekday() >= 5:
                nominal = 150000

            # Collect all final data
            claimable_data_dictionary={
                "date":tanggal.strftime('%d-%m-%Y'),
                "day":tanggal.strftime('%A'),
                "start":start_hour_str,
                "end":end_hour_str,
                "nominal":nominal,
                "activity":f"Customer : {customer}\nProject : {project}\nSO Number : {so}\nActivity : Onsite\n\n{activity_description}"
            }

            document_data_dictionary={
                "docx_date":tanggal.strftime('%d %B %Y'),
                "docx_image_path":f"{screenshot_folder}/{screenshot_naming}"
            }

            if nominal != 0:
                row.nth(index).screenshot(path=document_data_dictionary["docx_image_path"])
                csv_data_list.append(claimable_data_dictionary)
                document_data_list.append(document_data_dictionary)
                MODULE_LOGGER.info(f'Data number {index+1} is eligible for claim, check the csv for details.')
                print(f'Data number {index+1} is eligible for claim, check the csv for details.')
                print(f'The data is:\n{claimable_data_dictionary}\n----------\n')
            else:
                print(f'Data number {index+1} is not eligible for claim')
                MODULE_LOGGER.info(f'Data number {index+1} is not eligible for claim, check the console for details.')
                print(f'The data is:\n{claimable_data_dictionary}\nxxxxxxxxxx\n')
        else :
            MODULE_LOGGER.debug(f'Data number {index+1} / {row.count()} skipped due to filter')

    return {"csv_data":csv_data_list,"document_data":document_data_list}

def login_check(credential,page,page_number,browser_type):
    """
    Method to check if the page need to login.
    """
    MODULE_LOGGER.debug("Login check method called")
    if credential is None:
        try:
            with open('secret.json', 'r') as stream:
                data_json = json.load(stream)
            used_credential = data_json["ssportal"]
            MODULE_LOGGER.info("Credential succesfully obtained from json file")
        except:
            MODULE_LOGGER.error("Credential not provided via json file or cli ui")
            raise ImportError("""
            Credential is not provided and the json file is not found.
            Please provide the credential by creating json file based on the template.
            Or provide it on the CLI Prompter UI.
            """)
    else:
        used_credential = credential
        MODULE_LOGGER.info("Credential succesfully obtained from cli ui")

    if page_number is None or browser_type in ["HEADLESS","HEADED"]:
        while page.url != 'https://ssportal-tbs-2.packet-systems.com/activity':
            MODULE_LOGGER.info("Trying to login and navigate to site")
            MODULE_LOGGER.info(f"Currently at: {page.url}")
            page.goto('https://ssportal-tbs-2.packet-systems.com/activity')
            if page.url == 'https://ssportal-tbs-2.packet-systems.com/login':
                page.fill('id=username',used_credential["username"])
                page.fill('id=password',used_credential["password"])
                page.click('id=submitBtn')
            elif re.search("login\/error",page.url):
                MODULE_LOGGER.error("Something wrong while logging to the site. Retry the program.")
                raise EnvironmentError("""
                Either Username or Password is invalid while logging in!
                Please rerun the script.
                """)
    
    MODULE_LOGGER.info(f"Currently at: {page.url}")

def run(url_chrome,page_number,credential,browser_type):
    """
    TESTED AND CURRENTLY WORKING WITH THE UI PROMPT
    Main method to run the claim data scripts.
    """
    MODULE_LOGGER.info("SSPORTAL Claim collector script called")
    print("\nRunning SSPORTAL claim data collector scripts")

    path = Path(f"data/claim_data/{datetime.now().strftime('%d_%B_%Y')}")
    main_path = str(path)

    path.mkdir(parents=True, exist_ok=True)
    Path(f"{main_path}/screenshot").mkdir(parents=True, exist_ok=True)
    document_template = Document("data/template_capture.docx")
    csv_name = f"{main_path}/bipo_ssportal_claim_data-{datetime.now().strftime('%d-%m-%Y')}.csv"
    document_name = f"{main_path}/Capture SS Portal for Claim {datetime.now().strftime('%d %B %Y')}.docx"
    screenshot_folder = f"{main_path}/screenshot/{datetime.now().strftime('%d %B %Y')}"

    MODULE_LOGGER.info(f"Claim data will be saved at : {main_path}")
    with sync_playwright() as playwright:
        page = browser_initiator.initial_browser_checker(playwright,url_chrome,page_number,browser_type)
        page.set_default_timeout(60000)
        login_check(credential,page,page_number,browser_type)

        processed_data = claimable_data_collector(page,screenshot_folder)
        append_data_document(processed_data["document_data"],document_template)
        csv_processor.generate(processed_data["csv_data"],["date","nominal","activity","day","start","end"],csv_name)
        document_template.save(document_name)
        MODULE_LOGGER.info(f"CSV saved at : {csv_name}")
        MODULE_LOGGER.info(f"DOCX saved at : {document_name}")
        MODULE_LOGGER.info(f"Script finished at : {datetime.now().strftime('%H:%M %d-%m-%Y')}")
